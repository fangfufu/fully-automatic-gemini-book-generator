# main.py - Gemini Book Generator Core Script
# Current Version Date: 2025-04-19 (Based on last interaction)

# --- Imports ---
import os
import re
import time
import string
import datetime
import pickle
import hashlib
import tempfile
from dotenv import load_dotenv
from collections import OrderedDict

import google.generativeai as genai

# PDF Generation
from fpdf import FPDF

# DOCX Generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# PDF Merging
try:
    from pypdf import PdfWriter, PdfReader
except ImportError: print("ERROR: pypdf library not found. Run: pip install pypdf"); exit()

# Markdown Parsing
try:
    from markdown_it import MarkdownIt
except ImportError: print("ERROR: markdown-it-py not found. Run: pip install markdown-it-py"); exit()

# Config and Font Setup Imports
try: import config
except ImportError: print("ERROR: config.py not found."); exit()
try: from setup_fonts import setup_fonts, FONT_PATHS
except ImportError: print("ERROR: setup_fonts.py not found or missing FONT_PATHS."); exit()


# --- Configuration Loading & Constants ---
print("Loading environment variables from .env...")
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY: print("ERROR: GEMINI_API_KEY not found in .env file."); exit()

OUTPUT_DIR = "output"; CACHE_DIR = "api_cache"
PDF_WIDTH_MM = 152.4; PDF_HEIGHT_MM = 228.6; PDF_MARGIN_MM = 19.05
DOCX_WIDTH_INCHES = 6; DOCX_HEIGHT_INCHES = 9; DOCX_MARGIN_INCHES = 0.75
REGULAR_FONT_PATH = FONT_PATHS['regular']; BOLD_FONT_PATH = FONT_PATHS['bold']
ITALIC_FONT_PATH = FONT_PATHS['italic']; BOLD_ITALIC_FONT_PATH = FONT_PATHS['bold_italic']
DEFAULT_SECTION_WORDS_RANGE = (700, 1200); DEFAULT_TARGET_AUDIENCE = "General audience"
DEFAULT_BOOK_TYPE = "Non-Fiction"; DEFAULT_WORLD_SETTING = "Default real world"
DEFAULT_BOOK_STYLE = "Standard non-fiction"; DEFAULT_WRITING_STYLE = "Clear and informative"
DEFAULT_APPROX_CHAPTERS = 8; DEFAULT_APPROX_SECTIONS = 5

os.makedirs(OUTPUT_DIR, exist_ok=True); os.makedirs(CACHE_DIR, exist_ok=True)
print(f"Output directory set to: {os.path.abspath(OUTPUT_DIR)}")
print(f"API Cache directory set to: {os.path.abspath(CACHE_DIR)}")

# --- Markdown Parser Initialization ---
print("Initializing Markdown parser...")
md_parser = MarkdownIt()
print("Markdown parser initialized.")

# --- Gemini Setup ---
print("Configuring Gemini API...")
try:
    if not hasattr(config, 'GEMINI_MODEL') or not config.GEMINI_MODEL: raise AttributeError("GEMINI_MODEL")
    print(f"Using Gemini model specified in config: {config.GEMINI_MODEL}")
    genai.configure(api_key=API_KEY); model = genai.GenerativeModel(config.GEMINI_MODEL)
    print("Gemini API configured successfully.")
except AttributeError as e: print(f"ERROR: Config error: {e}"); exit()
except Exception as e: print(f"ERROR configuring Gemini: {e}"); exit()


# --- Helper Functions ---
def clean_filename(name):
    print(f"Cleaning filename for: '{name}'")
    valid_chars = "-_.() %s%s" % (string.ascii_letters, string.digits); cleaned_name = ''.join(c for c in name if c in valid_chars)
    cleaned_name = cleaned_name.replace(' ', '_'); cleaned_name = cleaned_name[:100]
    print(f"Cleaned filename: '{cleaned_name}'"); return cleaned_name if cleaned_name else "untitled_book"

# --- Gemini Generation Function with Caching ---
def generate_with_gemini(prompt_text, retries=3, delay=5):
    prompt_hash = hashlib.sha256(prompt_text.encode('utf-8')).hexdigest()
    cache_filename = f"{prompt_hash}.pkl"; cache_filepath = os.path.join(CACHE_DIR, cache_filename)
    if os.path.exists(cache_filepath):
        print(f"--- Cache HIT for prompt hash: {prompt_hash[:8]}...")
        try:
            with open(cache_filepath, 'rb') as f_cache: cached_data = pickle.load(f_cache)
            print(f"--- Loaded response from cache: {cache_filepath}")
            if isinstance(cached_data, str): return cached_data
            else: print(f"WARN: Cached data type mismatch. Miss.")
        except Exception as e: print(f"WARN: Failed loading cache file {cache_filepath}: {e}. Miss.")
    else: print(f"--- Cache MISS for prompt hash: {prompt_hash[:8]}...")
    print(f"\n--- Calling Gemini API (prompt length: {len(prompt_text)} chars)...")
    api_result_text = None
    for attempt in range(retries):
        try:
            print(f"Attempt {attempt + 1}/{retries}...")
            response = model.generate_content(prompt_text)
            if hasattr(response, 'text'):
                 api_result_text = response.text.strip()
                 print(f"--- Gemini call successful (length: {len(api_result_text)} chars)")
                 if not api_result_text or len(api_result_text) < 20: print(f"WARN: Short/empty response. Retrying..."); api_result_text = None; time.sleep(delay); continue
                 break
            else: print(f"WARN: Could not access .text. Attempt {attempt+1}/{retries}"); time.sleep(delay); continue
        except Exception as e:
            print(f"ERROR during API call (Attempt {attempt + 1}/{retries}): {e}")
            if "API key not valid" in str(e): return None
            if attempt == retries - 1: break
            time.sleep(delay)
    if api_result_text:
        try:
            with open(cache_filepath, 'wb') as f_cache: pickle.dump(api_result_text, f_cache)
            print(f"--- Saved API result to cache: {cache_filepath}")
        except Exception as e: print(f"WARN: Failed to save result to cache: {e}")
        return api_result_text
    else: print(f"ERROR: Failed Gemini call after {retries} attempts."); return None

# --- Other Generation Functions ---
def generate_title(topic, target_audience, book_type, world_setting, book_style, writing_style):
    print(f"\n--- Generating Title (Audience: {target_audience}, Style: {book_style}, Writing Style: {writing_style}) ---")
    topic_desc = f"non-fiction topic: '{topic}'"; style_desc = f"a standard '{book_style}'"; context_desc = ""
    if book_type == 'Fiction->Textbook Style': topic_desc = f"fictional topic: '{topic}'"; style_desc = f"a '{book_style}' from within that universe"; context_desc = f"The book is set within the '{world_setting}' context."
    prompt = f"""Generate 5 potential book titles for a book exploring the {topic_desc}. {context_desc} Intended style: {style_desc}. Writing style: '{writing_style}'. Target audience: '{target_audience}'. Tone should be appropriate. Present as simple list, one title per line, no numbers/bullets/quotes."""
    titles_text = generate_with_gemini(prompt)
    if not titles_text: print("WARN: Failed title gen. Fallback."); return "Untitled Book"
    possible_titles = [t.strip() for t in titles_text.split('\n') if t.strip()]
    if not possible_titles: print("WARN: No titles found. Fallback."); return "Untitled Book"
    print(f"Generated Titles Choices:\n{titles_text}"); selected_title = possible_titles[0]; return selected_title

def generate_toc(title, topic, book_type, world_setting, book_style, approx_chapters, approx_sections_per_chapter):
    print(f"\n--- Generating ToC (Target: ~{approx_chapters} Ch, ~{approx_sections_per_chapter} Sec/Ch) ---")
    topic_desc = f"non-fiction subject '{topic}'"; style_guidance = f"a standard '{book_style}'"; context_desc = ""
    structure_example = "Ex non-fiction:\n1. Chap Title\n  1.1. Section Title\n  1.2. Section Title\n2. Chap Title\n  2.1. Section Title" # Added dot to example
    if book_type == 'Fiction->Textbook Style': topic_desc = f"fictional subject '{topic}'"; style_guidance = f"a '{book_style}' from '{world_setting}' universe"; context_desc = f"Context: '{world_setting}'."; structure_example = f"Ex in-universe:\n1. Chapter\n  1.1. Section\n  1.2. Section\n2. Chapter\n  2.1. Section" # Added dot to example
    prompt = f"""Generate a detailed two-level Table of Contents for book '{title}' about {topic_desc}. {context_desc} Style: {style_guidance}. Logical structure, include Intro/Conclusion. Aim ~{approx_chapters} chapters. Aim ~{approx_sections_per_chapter} sections/chapter. Prioritize flow. {structure_example}\n\nFormat *exactly* like examples: Chapters start 'number.', sections start '  number.number.'. Use precise indentation AND include the final period after the section number.\n\nList *only* titles in this format. No other text.""" # Emphasized final period
    toc_text = generate_with_gemini(prompt); return toc_text

# --- Updated ToC Parser ---
def parse_toc(toc_text):
    """Parses two-level ToC into OrderedDict {chapter_title: [section_title1, ...]}. Returns None on failure."""
    print("\n--- Parsing Two-Level Table of Contents ---")
    if not toc_text: print("ERROR: Cannot parse empty ToC text."); return None
    print("--- Raw ToC Text ---"); print(toc_text); print("--------------------")
    parsed_toc = OrderedDict(); current_chapter_title = None
    # Regex for chapter: optional leading space, number, dot, mandatory space, title
    chapter_pattern = re.compile(r"^\s*(\d+)\.\s+(.+)")
    # Regex for section: mandatory leading space, number.number, DOT, mandatory space, title
    section_pattern = re.compile(r"^\s+(\d+\.\d+)\.\s+(.+)") # CORRECTED REGEX

    lines = toc_text.strip().split('\n')
    for line in lines:
        line = line.rstrip(); # Keep leading space for regex matching
        chapter_match = chapter_pattern.match(line)
        section_match = section_pattern.match(line)

        if chapter_match:
            chapter_num_str, chapter_title = chapter_match.groups(); chapter_title = chapter_title.strip()
            # Check if title is empty after stripping potential trailing spaces/chars
            if not chapter_title: print(f"WARN: Skipping chapter line with empty title: '{line}'"); continue
            current_chapter_title = chapter_title; parsed_toc[current_chapter_title] = []
            print(f"Found Chapter: {chapter_num_str}. {current_chapter_title}")
        elif section_match and current_chapter_title:
            section_num_str, section_title = section_match.groups(); section_title = section_title.strip()
            # Check if title is empty
            if not section_title: print(f"WARN: Skipping section line with empty title: '{line}'"); continue
            if current_chapter_title in parsed_toc:
                 parsed_toc[current_chapter_title].append(section_title);
                 print(f"  Found Section: {section_num_str}. {section_title} (under '{current_chapter_title}')")
            else: # Should not happen if chapter was matched first, but safety check
                 print(f"WARN: Section '{section_title}' found but current chapter '{current_chapter_title}' not in dict. Skipping.")
        elif line.strip(): # Check if line has content before printing warning
            print(f"WARN: Skipping non-matching ToC line: '{line}'")

    if not parsed_toc: print("ERROR: Failed to parse any chapters."); return None # Changed error message slightly
    print("\n--- Parsed ToC Structure ---"); total_sections = sum(len(s) for s in parsed_toc.values())
    # Check if any chapters ended up with no sections
    chapters_with_no_sections = [ch for ch, secs in parsed_toc.items() if not secs]
    if chapters_with_no_sections:
        print(f"WARN: The following chapters have no sections parsed: {', '.join(chapters_with_no_sections)}")
        # Decide if this is fatal? For now, proceed but warn heavily.
        if total_sections == 0 and len(parsed_toc) > 0:
             print("ERROR: No sections parsed for any chapter. Cannot proceed with content generation.")
             return None

    print(f"Parsed {len(parsed_toc)} chapters and {total_sections} sections."); return parsed_toc


def generate_section_content(section_title, chapter_title, book_title, topic, word_range, target_audience, book_type, world_setting, book_style, writing_style):
    # (No change needed)
    print(f"\n--- Generating Content for Section: '{section_title}' (Chapter: '{chapter_title}', Target words: {word_range[0]}-{word_range[1]}) ---")
    style_persona_instruction = ""
    if book_type == 'Fiction->Textbook Style': style_persona_instruction = f"You are an author writing an authoritative '{book_style}' titled '{book_title}' within universe '{world_setting}' about '{topic}' treated as real."
    else: style_persona_instruction = f"You are an author writing a non-fiction book titled '{book_title}' about '{topic}'. The book's style is '{book_style}'."
    prompt = f"""{style_persona_instruction} Target audience: '{target_audience}'. Write content for section '{section_title}' in chapter '{chapter_title}'. Focus *only* on this section, assuming chapter context. Aim {word_range[0]}-{word_range[1]} words. Writing style: '{writing_style}', tailored for audience/book style. Adapt complexity/tone. Address section title theme. Structure logically with paragraphs. No markdown formatting (like **bold** or *italic*). **Important:** Output *only* body text for section '{section_title}'. No titles or intro/concluding remarks."""
    content = generate_with_gemini(prompt, retries=3, delay=5)
    if not content: print(f"WARN: Failed generation for section '{section_title}'. Placeholder."); return f"(Content generation failed: {section_title})"
    actual_words = len(content.split()); print(f"Generated section word count: {actual_words} (Target: {word_range[0]}-{word_range[1]})")
    if actual_words < word_range[0] * 0.8: print(f"WARN: Section '{section_title}' shorter than target.")
    elif actual_words > word_range[1] * 1.2: print(f"WARN: Section '{section_title}' longer than target.")
    return content


# --- PDF Output Class ---
# (No change)
class PDF(FPDF):
    def header(self): pass
    def footer(self):
        current_page = self.page_no();
        if current_page > 2:
            self.set_y(-15);
            try: self.set_font('DejaVu', '', 8)
            except RuntimeError: self.set_font('Arial', '', 8)
            display_page_num = current_page - 2; self.cell(0, 10, f'{display_page_num}', 0, 0, 'C')


# --- Multi-Pass PDF Creation Function ---
# (No change needed here - markdown handling and structure already implemented)
def create_pdf(title, parsed_toc, all_content, filename, copyright_holder):
    """Creates PDF using multi-pass for page numbers, handling basic markdown."""
    print(f"\n--- Creating PDF with Sections & Basic Markdown: {filename} ---")
    content_pdf = None; toc_pdf = None; content_filename = None; toc_filename = None
    writer = PdfWriter()
    try:
        # Pass 1: Content PDF & Page Mapping
        print("--- Pass 1: Generating Content Pages ---")
        pdf_content = PDF(orientation='P', unit='mm', format=(PDF_WIDTH_MM, PDF_HEIGHT_MM))
        pdf_content.set_auto_page_break(auto=True, margin=PDF_MARGIN_MM)
        pdf_content.set_margins(left=PDF_MARGIN_MM, top=PDF_MARGIN_MM, right=PDF_MARGIN_MM)
        fonts_available = True
        try:
            pdf_content.add_font('DejaVu', '', REGULAR_FONT_PATH, uni=True); pdf_content.add_font('DejaVu', 'B', BOLD_FONT_PATH, uni=True)
            pdf_content.add_font('DejaVu', 'I', ITALIC_FONT_PATH, uni=True); pdf_content.add_font('DejaVu', 'BI', BOLD_ITALIC_FONT_PATH, uni=True)
        except RuntimeError as e: print(f"ERROR: PDF Pass 1 font missing: {e}"); fonts_available = False
        if not fonts_available: return False
        page_map = {}

        # Title Page
        pdf_content.add_page(); page_map[('TITLE', None)] = pdf_content.page_no()
        pdf_content.set_font('DejaVu', 'B', 24); pdf_content.ln(PDF_HEIGHT_MM / 4); pdf_content.multi_cell(0, 15, title, align='C')
        # Copyright Page
        pdf_content.add_page(); page_map[('COPYRIGHT', None)] = pdf_content.page_no()
        current_year = datetime.date.today().year; pdf_content.set_font('DejaVu', '', 10); pdf_content.set_y(PDF_HEIGHT_MM * 0.4)
        copyright_text = f"Copyright © {current_year} by {copyright_holder}\nAll rights reserved...\nDisclaimer: ... Reader discretion is advised." # Abridged
        pdf_content.multi_cell(0, 5, copyright_text, align='C')
        # Chapters and Sections Content
        print("Generating main content with basic markdown parsing and mapping pages...")
        for chapter_title, section_list in parsed_toc.items():
            pdf_content.add_page(); current_page = pdf_content.page_no(); page_map[(chapter_title, None)] = current_page
            pdf_content.set_font('DejaVu', 'B', 16); pdf_content.multi_cell(0, 10, f"{chapter_title}", ln=True, align='L'); pdf_content.ln(8)
            for section_title in section_list:
                current_page = pdf_content.page_no(); page_map[(chapter_title, section_title)] = current_page
                pdf_content.set_font('DejaVu', 'B', 13); pdf_content.multi_cell(0, 8, f"{section_title}", ln=True, align='L'); pdf_content.ln(4)
                section_content_text = all_content.get(chapter_title, {}).get(section_title, "Error: Content missing.")
                pdf_content.set_font('DejaVu', '', 11); is_bold = False; is_italic = False
                def set_pdf_style(pdf, bold, italic): style_str = ('B' if bold else '') + ('I' if italic else ''); pdf.set_font('DejaVu', style_str, 11)
                tokens = md_parser.parse(section_content_text)
                for token in tokens: # Process markdown tokens
                    if token.type == 'paragraph_open': pass
                    elif token.type == 'inline' and token.children:
                        for child in token.children:
                            if child.type == 'text': set_pdf_style(pdf_content, is_bold, is_italic); pdf_content.write(7, child.content)
                            elif child.type == 'strong_open': is_bold = True
                            elif child.type == 'strong_close': is_bold = False
                            elif child.type == 'em_open': is_italic = True
                            elif child.type == 'em_close': is_italic = False
                            elif child.type == 'softbreak' or child.type == 'hardbreak': pdf_content.ln(7)
                    elif token.type == 'paragraph_close': pdf_content.ln(7+2)
                pdf_content.ln(6) # Space after section
        # Save content temp file
        content_file_handle, content_filename = tempfile.mkstemp(suffix=".pdf", prefix="content_"); os.close(content_file_handle)
        pdf_content.output(content_filename); print(f"Content pages generated ({pdf_content.page_no()} pages). Temp file: {content_filename}")

        # === Pass 2: Generate ToC PDF ===
        print("--- Pass 2: Generating ToC Page(s) ---")
        pdf_toc = PDF(orientation='P', unit='mm', format=(PDF_WIDTH_MM, PDF_HEIGHT_MM))
        pdf_toc.set_auto_page_break(auto=True, margin=PDF_MARGIN_MM); pdf_toc.set_margins(left=PDF_MARGIN_MM, top=PDF_MARGIN_MM, right=PDF_MARGIN_MM)
        if not (os.path.exists(REGULAR_FONT_PATH) and os.path.exists(BOLD_FONT_PATH)): print(f"ERROR: PDF Pass 2 fonts failed."); return False
        pdf_toc.add_font('DejaVu', '', REGULAR_FONT_PATH, uni=True); pdf_toc.add_font('DejaVu', 'B', BOLD_FONT_PATH, uni=True)
        pdf_toc.add_page(); pdf_toc.set_font('DejaVu', 'B', 16); pdf_toc.cell(0, 10, "Table of Contents", ln=True, align='C'); pdf_toc.ln(10)
        chapter_num = 0
        for chapter_title, section_list in parsed_toc.items():
            chapter_num += 1; pdf_toc.set_font('DejaVu', 'B', 12)
            ch_page = page_map.get((chapter_title, None), 0); display_ch_page = max(0, ch_page - 2)
            toc_line = f"{chapter_num}. {chapter_title}"
            pdf_toc.multi_cell(PDF_WIDTH_MM - PDF_MARGIN_MM*2 - 15, 8, toc_line, ln=0); pdf_toc.cell(15, 8, str(display_ch_page) if display_ch_page > 0 else '', ln=1, align='R')
            pdf_toc.set_font('DejaVu', '', 11)
            section_num = 0
            for section_title in section_list:
                section_num += 1; pdf_toc.set_x(PDF_MARGIN_MM + 5)
                sec_page = page_map.get((chapter_title, section_title), 0); display_sec_page = max(0, sec_page - 2)
                available_width = PDF_WIDTH_MM - PDF_MARGIN_MM * 2 - 5
                toc_sec_line = f"  {chapter_num}.{section_num} {section_title}"
                pdf_toc.multi_cell(available_width - 15, 7, toc_sec_line, ln=0); pdf_toc.cell(15, 7, str(display_sec_page) if display_sec_page > 0 else '', ln=1, align='R')
            pdf_toc.ln(2)
        toc_file_handle, toc_filename = tempfile.mkstemp(suffix=".pdf", prefix="toc_"); os.close(toc_file_handle)
        pdf_toc.output(toc_filename); print(f"ToC page(s) generated. Temp file: {toc_filename}")

        # === Pass 3: Merge PDFs ===
        print("--- Pass 3: Merging PDF Files ---")
        content_reader = PdfReader(content_filename); toc_reader = PdfReader(toc_filename)
        writer.add_page(content_reader.pages[0]); writer.add_page(content_reader.pages[1])
        for page in toc_reader.pages: writer.add_page(page)
        if len(content_reader.pages) > 2:
             for i in range(2, len(content_reader.pages)): writer.add_page(content_reader.pages[i])
        with open(filename, "wb") as f_out: writer.write(f_out)
        print(f"--- PDF successfully merged and created: {filename} ---"); return True

    except Exception as e: print(f"ERROR during multi-pass PDF creation: {e}"); import traceback; traceback.print_exc(); return False
    finally:
        # --- Cleanup Temporary Files ---
        print("--- Cleaning up temporary PDF files ---")
        if content_filename and os.path.exists(content_filename):
            try:
                os.remove(content_filename)
                print(f"Cleaned up temp content file: {content_filename}")
            except OSError as e:
                print(f"WARN: Could not delete temp file {content_filename}: {e}")
        if toc_filename and os.path.exists(toc_filename):
            try:
                os.remove(toc_filename)
                print(f"Cleaned up temp ToC file: {toc_filename}")
            except OSError as e:
                print(f"WARN: Could not delete temp file {toc_filename}: {e}")

# --- DOCX Output Function ---
# (No change needed)
def create_docx(title, parsed_toc, all_content, filename, copyright_holder):
    print(f"\n--- Creating DOCX with Sections & Basic Markdown: {filename} ---")
    try:
        doc = Document(); current_year = datetime.date.today().year; section = doc.sections[0]
        section.page_width = Inches(DOCX_WIDTH_INCHES); section.page_height = Inches(DOCX_HEIGHT_INCHES)
        for M in ['left_margin', 'right_margin', 'top_margin', 'bottom_margin']: setattr(section, M, Inches(DOCX_MARGIN_INCHES))
        doc.add_heading(title, level=0)
        doc.add_page_break()
        def add_centered_para(text, pt_size=10, is_italic=False): p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; run = p.add_run(text); run.font.size = Pt(pt_size); run.italic = is_italic; return p
        add_centered_para(f"Copyright © {current_year} by {copyright_holder}"); add_centered_para("All rights reserved."); doc.add_paragraph()
        add_centered_para(f"Published in Cambridge, UK"); add_centered_para(f"First Edition, {current_year}")
        doc.add_paragraph(); add_centered_para("Disclaimer: ... Reader discretion is advised.", pt_size=9, is_italic=True)
        doc.add_page_break(); doc.add_heading("Table of Contents", level=1)
        p_toc_instr = doc.add_paragraph(); run_toc_instr = p_toc_instr.add_run("Note: Update this ToC field in Word...\nInstructions: Right-click -> 'Update Field' -> 'Update entire table'. (Uses Heading 1/2 styles).")
        run_toc_instr.italic = True; run_toc_instr.font.size = Pt(10)
        print("Adding chapter (H1) and section (H2) content to DOCX with basic markdown parsing...")
        for chapter_title, section_list in parsed_toc.items():
            doc.add_page_break(); doc.add_heading(chapter_title, level=1)
            for section_title in section_list:
                section_content_text = all_content.get(chapter_title, {}).get(section_title, "Error: Content not found.")
                doc.add_heading(section_title, level=2)
                tokens = md_parser.parse(section_content_text); current_paragraph = None; is_bold = False; is_italic = False
                for token in tokens:
                    if token.type == 'paragraph_open': current_paragraph = doc.add_paragraph(); is_bold = False; is_italic = False
                    elif token.type == 'inline' and current_paragraph is not None and token.children:
                        for child in token.children:
                            if child.type == 'text': run = current_paragraph.add_run(child.content); run.bold = is_bold; run.italic = is_italic
                            elif child.type == 'strong_open': is_bold = True
                            elif child.type == 'strong_close': is_bold = False
                            elif child.type == 'em_open': is_italic = True
                            elif child.type == 'em_close': is_italic = False
                    elif token.type == 'paragraph_close': current_paragraph = None; is_bold = False; is_italic = False
        doc.save(filename); print(f"--- DOCX successfully created: {filename} ---"); print("      Action Required: Open DOCX and update ToC field."); return True
    except Exception as e: print(f"ERROR creating DOCX: {e}"); import traceback; traceback.print_exc(); return False


# --- Main Execution Block ---
if __name__ == "__main__":
    print("\n========= Fully Automatic Gemini Book Generator =========")
    start_time = time.time()
    # --- Validate Configuration ---
    print("\n--- Validating Configuration from config.py ---")
    try:
        # (Validation logic remains the same)
        if not hasattr(config, 'BOOK_TOPIC') or not isinstance(config.BOOK_TOPIC, str): raise AttributeError("BOOK_TOPIC")
        book_topic_from_config = config.BOOK_TOPIC.strip(); assert book_topic_from_config, "BOOK_TOPIC empty."
        book_type_from_config = getattr(config, 'BOOK_TYPE', DEFAULT_BOOK_TYPE)
        if book_type_from_config not in ['Non-Fiction', 'Fiction->Textbook Style']: print(f"WARN: Invalid BOOK_TYPE."); book_type_from_config = DEFAULT_BOOK_TYPE
        world_setting_from_config = DEFAULT_WORLD_SETTING
        if book_type_from_config == 'Fiction->Textbook Style':
             if not hasattr(config, 'WORLD_SETTING') or not isinstance(config.WORLD_SETTING, str): raise AttributeError("WORLD_SETTING required")
             world_setting_from_config = config.WORLD_SETTING.strip(); assert world_setting_from_config, "WORLD_SETTING empty"
        book_style_from_config = getattr(config, 'BOOK_STYLE', DEFAULT_BOOK_STYLE).strip()
        if not book_style_from_config: print(f"WARN: BOOK_STYLE empty."); book_style_from_config = DEFAULT_BOOK_STYLE
        writing_style_from_config = getattr(config, 'WRITING_STYLE', DEFAULT_WRITING_STYLE).strip()
        if not writing_style_from_config: print(f"WARN: WRITING_STYLE empty."); writing_style_from_config = DEFAULT_WRITING_STYLE
        copyright_holder_from_config = getattr(config, 'COPYRIGHT_HOLDER', '').strip()
        if not copyright_holder_from_config or copyright_holder_from_config == "Your Name or Pseudonym Here": print("WARN: COPYRIGHT_HOLDER default/empty."); copyright_holder_from_config = "[Your Name Here - Edit in config.py]"
        target_section_words_range = getattr(config, 'TARGET_SECTION_WORDS_RANGE', DEFAULT_SECTION_WORDS_RANGE)
        if not isinstance(target_section_words_range, (tuple, list)) or len(target_section_words_range) != 2 or not all(isinstance(n, int) for n in target_section_words_range) or target_section_words_range[0] < 50 or target_section_words_range[1] < target_section_words_range[0]: print(f"WARN: TARGET_SECTION_WORDS_RANGE invalid."); target_section_words_range = DEFAULT_SECTION_WORDS_RANGE
        target_audience_from_config = getattr(config, 'TARGET_AUDIENCE', DEFAULT_TARGET_AUDIENCE).strip()
        if not target_audience_from_config: print(f"WARN: TARGET_AUDIENCE empty."); target_audience_from_config = DEFAULT_TARGET_AUDIENCE
        approx_chapters_from_config = getattr(config, 'APPROX_CHAPTERS', DEFAULT_APPROX_CHAPTERS)
        if not isinstance(approx_chapters_from_config, int) or approx_chapters_from_config < 3: print(f"WARN: APPROX_CHAPTERS invalid."); approx_chapters_from_config = DEFAULT_APPROX_CHAPTERS
        approx_sections_per_chapter_from_config = getattr(config, 'APPROX_SECTIONS_PER_CHAPTER', DEFAULT_APPROX_SECTIONS)
        if not isinstance(approx_sections_per_chapter_from_config, int) or approx_sections_per_chapter_from_config < 1: print(f"WARN: APPROX_SECTIONS_PER_CHAPTER invalid."); approx_sections_per_chapter_from_config = DEFAULT_APPROX_SECTIONS

        print(f"\n--- Configuration Loaded ---")
        print(f" Book Topic: '{book_topic_from_config}' ({book_type_from_config} / {book_style_from_config})")
        print(f" Target Audience: '{target_audience_from_config}' / Style: '{writing_style_from_config}'")
        print(f" Structure Target: ~{approx_chapters_from_config} Chaps, ~{approx_sections_per_chapter_from_config} Secs/Chap")
        print(f" Section Length Target: {target_section_words_range[0]}-{target_section_words_range[1]} words")
        print("--------------------------")
    except (AttributeError, ValueError, AssertionError) as e: print(f"ERROR: Invalid config: {e}"); exit()
    except Exception as e: print(f"ERROR reading config: {e}"); exit()

    # --- Ensure fonts are available ---
    print("\n--- Checking Font Availability ---")
    fonts_ready = setup_fonts()
    if not fonts_ready: print("\nCRITICAL ERROR: Fonts setup failed. Exiting."); exit()
    print("--- Font setup complete/verified ---")

    # --- Start Book Generation Workflow ---
    print("\n--- Starting Book Generation ---")
    book_title = generate_title(book_topic_from_config, target_audience_from_config, book_type_from_config, world_setting_from_config, book_style_from_config, writing_style_from_config)
    if book_title == "Untitled Book": print("WARN: Using fallback title.")
    print(f"Selected Title: {book_title}")
    safe_title = clean_filename(book_title); pdf_filename = os.path.join(OUTPUT_DIR, f"{safe_title}_6x9_print.pdf"); docx_filename = os.path.join(OUTPUT_DIR, f"{safe_title}_ebook.docx")
    toc_raw_text = generate_toc(book_title, book_topic_from_config, book_type_from_config, world_setting_from_config, book_style_from_config, approx_chapters_from_config, approx_sections_per_chapter_from_config)
    if not toc_raw_text: print("ERROR: Failed to generate ToC. Exiting."); exit()
    parsed_toc_structure = parse_toc(toc_raw_text)
    if not parsed_toc_structure: print("ERROR: Failed to parse ToC structure. Exiting."); exit()

    print(f"\n--- Generating Content for Sections ---")
    all_sections_content = OrderedDict()
    total_sections_parsed = sum(len(secs) for secs in parsed_toc_structure.values()); current_section_count = 0
    print(f"Found {len(parsed_toc_structure)} chapters and {total_sections_parsed} sections in parsed ToC.")
    for chapter_title, section_list in parsed_toc_structure.items():
        all_sections_content[chapter_title] = OrderedDict()
        if not section_list: print(f"WARN: No sections for chapter '{chapter_title}'. Skipping."); continue
        print(f"\nProcessing Chapter: '{chapter_title}'")
        for section_title in section_list:
            current_section_count += 1
            print(f"\n>>> Processing Section {current_section_count}/{total_sections_parsed}: '{section_title}'")
            content = generate_section_content(
                section_title, chapter_title, book_title, book_topic_from_config, target_section_words_range,
                target_audience_from_config, book_type_from_config, world_setting_from_config, book_style_from_config, writing_style_from_config
            )
            all_sections_content[chapter_title][section_title] = content

    print("\n--- Generating Final Output Files ---")
    pdf_success = create_pdf(book_title, parsed_toc_structure, all_sections_content, pdf_filename, copyright_holder_from_config)
    docx_success = create_docx(book_title, parsed_toc_structure, all_sections_content, docx_filename, copyright_holder_from_config)

    end_time = time.time(); duration = end_time - start_time; completion_time_struct = time.localtime(); completion_time_str = time.strftime("%Y-%m-%d %H:%M:%S %Z", completion_time_struct)
    actual_chapters = len(parsed_toc_structure) if parsed_toc_structure else 0; actual_sections = total_sections_parsed if parsed_toc_structure else 0
    print("\n========= Generation Summary =========")
    print(f"Book Title: {book_title}"); print(f"Topic: {book_topic_from_config}")
    print(f"Actual Structure: {actual_chapters} Chapters / {actual_sections} Sections")
    if pdf_success: print(f"PDF (Multi-pass, Basic Markdown) saved to: {os.path.abspath(pdf_filename)}")
    else: print("PDF generation failed or skipped.")
    if docx_success: print(f"DOCX (Ebook, Basic Markdown, Manual ToC) saved to: {os.path.abspath(docx_filename)}"); print("      -> Remember to update DOCX ToC field.")
    else: print("DOCX generation failed or skipped.")
    print(f"Total execution time: {duration:.2f} seconds"); print(f"Completed at: {completion_time_str}")
    print("====================================")